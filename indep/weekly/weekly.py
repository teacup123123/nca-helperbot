import dataclasses
import os
import shutil
from contextlib import contextmanager

from docx import Document as rDoc
from docx.document import Document
from docx.table import Table, _Row, _Cell
from datetime import datetime as Datetime, date as Date, time as Time

format13 = '請假表單周13', 'name time type hrs exempt_ln exempt_tn'
format24 = '請假表單周24', 'name time type hrs exempt_ln exempt_tn morning_sport'
format5 = '請假表單周5', 'name time type hrs exempt_ln'
weekdayformats = [format13, format24, format13, format24, format5]

here, _ = os.path.split(__file__)

morning = [Time(8, 45), Time(9, 45), Time(10, 45), Time(11, 45)]
afternoon = [Time(13, 45), Time(14, 45), Time(15, 45), Time(16, 45)]
allday = morning + afternoon


# morningmsk = sum(1 << i for i in range(4))
# afternoonmsk = sum(1 << i for i in range(4, 8))
# fulldaymsk = sum(1 << i for i in range(8))


@dataclasses.dataclass
class OutputWeekday:
    out: str
    date: Date

    def add(self, name, dorm, holiday_type, iso_signature: str):
        song8: Table = self._d.tables[0]
        song9: Table = self._d.tables[1]
        holiday: Table = self._d.tables[2]
        starttime, endtime = map(Time.fromisoformat, iso_signature.split('-'))

        if dorm == '松八':
            table = song8
        elif dorm == '松九':
            table = song9
        else:
            raise ValueError
        table.add_row()
        rdorm: _Row = table.rows[-1]
        #
        # for rdorm in table.rows:
        #
        # while True:
        #     for r in r
        #     try:
        #         rdorm: _Row = next(self._it8)
        #     except StopIteration:
        #         break
        # if dorm == '松八':
        #     rdorm: _Row = next(self._it8)
        # elif dorm == '松九':
        #     rdorm: _Row = next(self._it9)
        # else:
        #     raise ValueError

        print(name, holiday_type, dorm, iso_signature)
        if holiday_type == '預': holiday_type = '榮'
        holiday_type = holiday_type.replace('預', '')
        holiday_type = holiday_type.replace('+', '')
        if '補' in holiday_type: holiday_type = '補'
        assert len(holiday_type)
        found = False
        for hrow in holiday.rows:
            header = hrow.cells[0]
            # print(header.text)
            if holiday_type in header.text:
                found = True
                break
        assert found

        for ri, key in enumerate(self._format.split()):
            c = rdorm.cells[ri]
            c: _Cell
            if key == 'name':
                c.paragraphs[-1].text = name
            elif key == 'time':
                c.paragraphs[-1].text = iso_signature
            elif key == 'hrs':
                c.paragraphs[-1].text = f'{sum(starttime < w < endtime for w in allday)}'
            elif key == 'type':
                c.paragraphs[-1].text = holiday_type
            elif key == 'exempt_tn' or key == 'morning_sport':
                c.paragraphs[-1].text = '免' if all(starttime < w < endtime for w in morning) else '需'
            elif key == 'exempt_ln':
                c.paragraphs[-1].text = '免' if all(starttime < w < endtime for w in afternoon) else '需'

        if all(starttime < w < endtime for w in allday):
            hrow.cells[3].text = '\n'.join(hrow.cells[3].text.splitlines() + [name])
        elif all(starttime < w < endtime for w in afternoon):
            hrow.cells[2].text = '\n'.join(hrow.cells[2].text.splitlines() + [name])
        elif all(starttime < w < endtime for w in morning):
            hrow.cells[1].text = '\n'.join(hrow.cells[1].text.splitlines() + [name])
        else:
            hrow.cells[4].text = '\n'.join(hrow.cells[4].text.splitlines() + [name])

    def open(self):
        self._src, self._format = weekdayformats[self.date.weekday()]
        self._src = os.path.join(here, f'{self._src}.docx')
        shutil.copy(self._src, self.out)
        self._d: Document = rDoc(self.out)

        song8: Table = self._d.tables[0]
        song9: Table = self._d.tables[1]
        holiday: Table = self._d.tables[2]

        self._it8 = iter(song8.rows)
        self._it9 = iter(song9.rows)

        next(self._it8)
        next(self._it9)

        doc = self._d
        date = self.date
        title = (f' 日期：{date.year - 1911}-{date.month:02}-{date.day:02} '
                 f'(週{"一二三四五"[date.weekday()]})')
        doc.paragraphs[0].runs[-1].add_text(title)

    def close(self):
        self._d.save(self.out)
        pass


@contextmanager
def prepare(date: Date):
    out: OutputWeekday = None
    try:
        out = OutputWeekday(
            os.path.join(here, 'output', '請假表單{}週{}.docx'.format(
                *date.strftime(f'%Y%m%d {date.weekday() + 1}').split())),
            date
        )
        out.open()
        yield out
    finally:
        out.close()


if __name__ == '__main__':
    with prepare(Date.today()) as f:
        f: OutputWeekday
        f.add('miao', '松八', '榮', 255)
